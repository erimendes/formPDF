from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading('TERMO DE REFERÊNCIA', 0)

# 1. IDENTIFICAÇÃO DO DOCUMENTO
doc.add_heading('1. IDENTIFICAÇÃO DO DOCUMENTO', level=1)
doc.add_paragraph(
    '• Título: Termo de Referência para Aquisição de Licenças e/ou Contratação de Serviços de Suporte para Red Hat Enterprise Linux Server 7.9\n'
    '• Data de Elaboração: [Data Atual]\n'
    '• Órgão/Setor Responsável: [Nome do Órgão/Setor]\n'
    '• Responsável pela Elaboração: [Nome e Cargo do Responsável]\n'
)

# 2. OBJETO
doc.add_heading('2. OBJETO', level=1)

# 2.1 Descrição do Objeto (adjusted)
doc.add_heading('2.1 Descrição do Objeto', level=2)
doc.add_paragraph(
    'Aquisição de licenças de uso e/ou contratação de serviços de suporte técnico especializado para o sistema operacional Red Hat Enterprise Linux Server 7.9, '
    'com o objetivo de garantir a estabilidade e a segurança da infraestrutura de servidores da [Nome da Organização/Setor].'
)

# 2.2 Justificativa
doc.add_heading('2.2 Justificativa', level=2)
doc.add_paragraph(
    'O ambiente atual apresenta falhas na comunicação entre a console QRadar e o coletor de eventos, impossibilitando a aplicação de atualizações e regras. Dessa forma, '
    'torna-se necessária uma nova instalação do coletor, demandando licenças atualizadas e suporte técnico adequado para o Red Hat Enterprise Linux Server 7.9.\n\n'
    'A escolha pela versão Red Hat Enterprise Linux Server 7.9 justifica-se pela necessidade de compatibilidade com a solução de segurança atualmente adotada na organização, '
    'especificamente o IBM QRadar versão 7.5. De acordo com a documentação técnica da IBM, essa versão do QRadar é oficialmente suportada exclusivamente no RHEL 7.9.\n\n'
    'Atualizações para versões superiores do sistema operacional (como RHEL 8 ou 9) não são recomendadas, pois não são compatíveis com a versão atual do QRadar, o que pode '
    'comprometer a integridade, a estabilidade e a segurança do ambiente de coleta e análise de eventos.\n\n'
    'Adicionalmente, existem limitações técnicas quanto à atualização in-place de sistemas que utilizam partições criptografadas com LUKS, o que reforça a decisão pela manutenção '
    'da versão 7.9 como requisito técnico essencial para garantir o funcionamento adequado da solução de SIEM.'
)

# 2.3 Considerações Técnicas Complementares (Opcional)
doc.add_heading('2.3 Considerações Técnicas Complementares (Opcional)', level=2)
doc.add_paragraph(
    'Ciclo de Vida do RHEL 7.9: Ainda que a compatibilidade com o QRadar 7.5 seja o principal fator determinante para a escolha do Red Hat Enterprise Linux 7.9, é importante '
    'destacar que esta versão do sistema operacional atingiu o fim do suporte estendido em junho de 2024, conforme o ciclo de vida oficial da Red Hat. A decisão de mantê-la neste '
    'cenário específico foi tomada com base na necessidade de garantir a funcionalidade plena e o suporte da solução IBM QRadar, que depende dessa versão do sistema operacional.\n\n'
    'Estratégia de Migração Futura: A organização está ciente das limitações do ambiente atual e, caso esteja previsto em seu planejamento estratégico, poderá avaliar futuramente '
    'a atualização da solução QRadar para versões mais recentes, possibilitando também a migração do sistema operacional para versões mais atuais do RHEL. Esta medida, no entanto, '
    'dependerá da evolução do suporte da IBM e da viabilidade técnica e orçamentária da instituição.'
)

# 3. ESPECIFICAÇÕES TÉCNICAS
doc.add_heading('3. ESPECIFICAÇÕES TÉCNICAS', level=1)

# 3.1 Sistema Operacional
doc.add_heading('3.1 Sistema Operacional', level=2)
doc.add_paragraph(
    '• Versão: Red Hat Enterprise Linux Server 7.9 (Binary DVD)\n'
    '• Arquitetura: [x86_64 ou outra]\n'
    '• Tipo de Licenciamento: [Por servidor, por socket, ou outro modelo aplicável]\n'
    '• Quantidade de Licenças: [Número exato]\n'
    '• Direitos de Uso: [Especificar, conforme contrato ou termos da Red Hat]\n'
    '• Modo FIPS: [Sim/Não]. Se "Sim", configurar conforme a documentação oficial da Red Hat e IBM QRadar.'    
)

# 3.2 Configuração de Instalação
doc.add_heading('3.2 Configuração de Instalação', level=2)
doc.add_paragraph(
    '• Particionamento:\n'
    '  - Utilização de LVM (Logical Volume Management);\n'
    '  - Definição de pontos de montagem e capacidades mínimas com base na documentação oficial: "Linux operating system partition properties for QRadar installations on your own system".\n'
    '  - Caso a criptografia de disco seja necessária, utilizar LUKS (Linux Unified Key Setup), especificando os requisitos.\n'
    '  - Atenção: Sistemas com partições criptografadas por LUKS podem apresentar limitações no processo de atualização (upgrade in-place) para versões superiores ao Red Hat Enterprise Linux 8.8. '
    'A Red Hat não garante compatibilidade plena de atualização automática nesses casos. Assim, recomenda-se avaliar previamente se a criptografia será um requisito essencial, especialmente se houver '
    'planos de migração futura para versões mais recentes do sistema operacional.\n'
    '• Seleção de Software: Instalação mínima (Minimal Install)\n'
    '• Idioma: Inglês (US)\n'
    '• Fuso Horário: [Exemplo: America/Sao_Paulo]\n'
    '• Hostname: Configuração obrigatória de FQDN (Fully Qualified Domain Name)'
)

# 3.3 Configuração de Rede
doc.add_heading('3.3 Configuração de Rede', level=2)
doc.add_paragraph(
    '• Endereçamento IP manual (IPv4 e/ou IPv6)\n'
    '• Netmask/Prefixo, Gateway e no mínimo dois servidores DNS\n'
    '• Habilitar conexão automática com a rede'
)

# 3.4 Pós-Instalação
doc.add_heading('3.4 Pós-Instalação', level=2)
doc.add_paragraph(
    '• Instruções detalhadas para desativação do SELinux (destacando implicações de segurança)'
)

# 3.5 Serviços de Suporte Técnico (Opcional, conforme contratação)
doc.add_heading('3.5 Serviços de Suporte Técnico (Opcional, conforme contratação)', level=2)
doc.add_paragraph(
    '• Nível de Suporte: [Básico/Intermediário/Avançado]\n'
    '• Forma de Atendimento: [Remoto/Presencial/Híbrido]\n'
    '• Cobertura: [24x7, horário comercial, etc.]\n'
    '• SLA: [Ex: 4h para resposta, 8h para solução]\n'
    '• Atualizações e Correções: Inclusas conforme políticas da Red Hat\n'
    '• Base de Conhecimento: Acesso incluído\n'
    '• Apoio à Migração/Implantação: [Sim/Não – Especificar escopo]'
)

# 3.6 Requisitos Adicionais
doc.add_heading('3.6 Requisitos Adicionais', level=2)
doc.add_paragraph(
    '• Documentação técnica completa\n'
    '• Treinamento para equipe interna (se necessário)\n'
    '• Ferramentas de gerenciamento compatíveis\n'
    '• Garantia de compatibilidade com IBM QRadar'
)

# 4. CRITÉRIOS DE ACEITAÇÃO
doc.add_heading('4. CRITÉRIOS DE ACEITAÇÃO', level=1)
doc.add_paragraph(
    '• Entrega e validação das licenças Red Hat Enterprise Linux Server 7.9 na versão Binary DVD\n'
    '• Comprovação da instalação conforme as especificações detalhadas\n'
    '• Relatórios de instalação e configuração\n'
    '• Suporte técnico funcionando conforme o SLA acordado'
)

# 5. PRAZOS
doc.add_heading('5. PRAZOS', level=1)
doc.add_paragraph(
    '• Entrega das Licenças: [Ex: até 15 dias corridos após assinatura do contrato]\n'
    '• Início dos Serviços de Suporte: [Data ou condição para início]\n'
    '• Duração do Contrato/Licenças: [Ex: 12 meses, renováveis]'
)

# 6. LOCAL DE ENTREGA E EXECUÇÃO
doc.add_heading('6. LOCAL DE ENTREGA E EXECUÇÃO', level=1)
doc.add_paragraph(
    '• Entrega das Licenças: [Local físico ou digital]\n'
    '• Execução dos Serviços de Suporte: [Presencial, remoto, híbrido – conforme aplicável]'
)

# 7. OBRIGAÇÕES DO CONTRATANTE
doc.add_heading('7. OBRIGAÇÕES DO CONTRATANTE', level=1)
doc.add_paragraph(
    '• Designar um Gestor do Contrato responsável pela coordenação, definição de prioridades, acompanhamento da execução, validação dos serviços e atesto das faturas.\n'
    '• Disponibilizar as informações e acessos necessários para execução do contrato.'
)

# 8. OBRIGAÇÕES DO CONTRATADO
doc.add_heading('8. OBRIGAÇÕES DO CONTRATADO', level=1)
doc.add_paragraph(
    '• Fornecer as licenças conforme especificado neste Termo\n'
    '• Disponibilizar documentação técnica detalhada dos processos realizados\n'
    '• Prestar os serviços de suporte técnico conforme SLA e níveis contratados\n'
    '• Garantir a compatibilidade e conformidade com as exigências técnicas mencionadas'
)

# Save the document
doc_path = "Termo_de_Referencia_RHEL_7.9.docx"
doc.save(doc_path)
print(f"Documento salvo em: {doc_path}")
